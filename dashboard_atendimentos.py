import streamlit as st
import pandas as pd
import plotly.express as px
from datetime import date
from io import BytesIO

# ===============================
# IMPORTAÇÃO SEGURA DO DOCX
# ===============================
try:
    from docx import Document
    from docx.shared import Pt
    DOCX_DISPONIVEL = True
except ModuleNotFoundError:
    DOCX_DISPONIVEL = False

# ===============================
# CONFIGURAÇÃO DA PÁGINA
# ===============================
title = 'Análises de Atendimentos - UPA 24H Dona Zulmira Soares'
st.set_page_config(page_title=title, layout='wide')
st.image('TESTEIRA PAINEL UPA1.png', width=120)
st.title(title)

# ===============================
# SIDEBAR
# ===============================
st.sidebar.image('TESTEIRA PAINEL UPA1.png', width=300)
st.sidebar.header('Filtros')
uploaded_files = st.sidebar.file_uploader(
    "Envie as planilhas de atendimentos", type=["xlsx"], accept_multiple_files=True
)

# ===============================
# FUNÇÃO PROCESSAR PLANILHA
# ===============================
def processar_planilha(file):
    df = pd.read_excel(file, skiprows=1, header=None)
    df = df.dropna(how='all')
    df.columns = [
        'CPF', 'Paciente', 'Data', 'Hora', 'Especialidade', 'Profissional',
        'Motivo Alta', 'Procedimento', 'Cid10', 'Prioridade'
    ]
    return df

# ===============================
# FUNÇÕES DE GRÁFICOS
# ===============================
def criar_grafico_barra(df, coluna, titulo, top_n=10):
    contagem = df[coluna].value_counts().reset_index()
    contagem.columns = [coluna, 'Quantidade']
    return px.bar(
        contagem.head(top_n),
        x=coluna,
        y='Quantidade',
        title=titulo,
        color='Quantidade',
        template='plotly_white'
    )

def criar_grafico_pizza(df, coluna, titulo, top_n=10):
    contagem = df[coluna].value_counts().reset_index()
    contagem.columns = [coluna, 'Quantidade']
    return px.pie(
        contagem.head(top_n),
        names=coluna,
        values='Quantidade',
        title=titulo
    )

# ===============================
# PROCESSAMENTO PRINCIPAL
# ===============================
if uploaded_files:
    dfs = [processar_planilha(f) for f in uploaded_files]
    df_final = pd.concat(dfs, ignore_index=True)
    df_final = df_final.fillna('NÃO INFORMADO')

    # ===============================
    # FILTROS (CORRIGIDO)
    # ===============================
    for col in df_final.columns:
        valores = df_final[col].astype(str).unique().tolist()
        valores.sort()
        filtro = st.sidebar.multiselect(f'Filtrar por {col}', valores)
        if filtro:
            df_final = df_final[df_final[col].astype(str).isin(filtro)]

    # ===============================
    # DATA / TURNO
    # ===============================
    df_final['Data Atendimento'] = pd.to_datetime(df_final['Data'], errors='coerce')
    df_final['Hora_num'] = pd.to_datetime(df_final['Hora'], errors='coerce').dt.hour

    def identificar_turno(h):
        if pd.isnull(h): return 'Indefinido'
        if 6 <= h < 12: return 'Manhã'
        if 12 <= h < 18: return 'Tarde'
        if 18 <= h < 24: return 'Noite'
        return 'Madrugada'

    df_final['Turno'] = df_final['Hora_num'].apply(identificar_turno)

    # ===============================
    # RESOLUTIVIDADE – SUS / PNAU
    # ===============================
    def classificar_resolutividade(motivo):
        m = str(motivo).lower()
        if any(x in m for x in ['alta', 'prescrição', 'observação', 'encerramento']):
            return 'Resolvido na UPA'
        if any(x in m for x in ['transfer', 'óbito', 'regulado']):
            return 'Não resolvido na UPA'
        return 'Indefinido'

    df_final['Resolutividade'] = df_final['Motivo Alta'].apply(classificar_resolutividade)

    total = len(df_final)
    taxa_resolucao = (df_final['Resolutividade'] == 'Resolvido na UPA').mean()

    df_final = df_final.sort_values(['CPF', 'Data Atendimento'])
    df_final['Retorno_72h'] = (
        df_final.groupby('CPF')['Data Atendimento']
        .diff().dt.total_seconds().div(3600).le(72)
    )
    taxa_retorno = df_final['Retorno_72h'].mean()

    amarelos = df_final[df_final['Prioridade'].str.contains('Amarelo', case=False)]
    taxa_amarelo = (
        (amarelos['Resolutividade'] == 'Resolvido na UPA').mean()
        if len(amarelos) > 0 else 0
    )

    perfil = df_final['Prioridade'].value_counts(normalize=True) * 100
    verde_azul = perfil.filter(like='Verde').sum() + perfil.filter(like='Azul').sum()

    score = taxa_resolucao * 0.4 + (1 - taxa_retorno) * 0.2 + taxa_amarelo * 0.4

    if score >= 0.80:
        status = '🟢 UPA RESOLUTIVA'
    elif score >= 0.60:
        status = '🟡 PARCIALMENTE RESOLUTIVA'
    else:
        status = '🔴 BAIXA RESOLUTIVIDADE'

    # ===============================
    # PAINEL GERENCIAL
    # ===============================
    st.markdown("## 🏥 Avaliação de Resolutividade – SUS")

    c1, c2, c3, c4 = st.columns(4)
    c1.metric("Resolução na UPA", f"{taxa_resolucao:.1%}")
    c2.metric("Retorno ≤ 72h", f"{taxa_retorno:.1%}")
    c3.metric("Resolução Amarelos", f"{taxa_amarelo:.1%}")
    c4.metric("Score Geral", f"{score:.2f}", status)

    if verde_azul > 60:
        st.warning(
            f"⚠️ {verde_azul:.1f}% dos atendimentos são Verde/Azul — "
            "indício de sobrecarga da Atenção Básica."
        )

    fig_res = px.histogram(
        df_final,
        x='Resolutividade',
        color='Prioridade',
        barmode='group',
        title='Desfecho dos Atendimentos por Classificação de Risco'
    )
    st.plotly_chart(fig_res, use_container_width=True)

    # ===============================
    # ANÁLISES EXISTENTES
    # ===============================
    colunas_para_analisar = [
        'Especialidade', 'Motivo Alta', 'Profissional',
        'Prioridade', 'Cid10', 'Procedimento'
    ]

    top_n = st.sidebar.slider("Número de itens no gráfico", 5, 20, 10)
    tipo_grafico = st.sidebar.selectbox("Tipo de gráfico", ["Barras", "Pizza"])

    for col in colunas_para_analisar:
        st.subheader(f"Análises para {col}")
        if tipo_grafico == "Barras":
            st.plotly_chart(
                criar_grafico_barra(df_final, col, f'Top {top_n} {col}', top_n),
                use_container_width=True
            )
        else:
            st.plotly_chart(
                criar_grafico_pizza(df_final, col, f'Top {top_n} {col}', top_n),
                use_container_width=True
            )

    # ===============================
    # RELATÓRIO DOCX – RAG / PAS / TCE
    # ===============================
    st.markdown("## 📄 Relatório Oficial")

    def gerar_relatorio_docx():
        doc = Document()
        doc.add_heading('Relatório de Avaliação da UPA 24h Dona Zulmira Soares', 1)

        p = doc.add_paragraph()
        p.add_run(f"Total de atendimentos: {total}\n")
        p.add_run(f"Taxa de resolução: {taxa_resolucao:.1%}\n")
        p.add_run(f"Retorno até 72h: {taxa_retorno:.1%}\n")
        p.add_run(f"Resolução casos amarelos: {taxa_amarelo:.1%}\n")
        p.add_run(f"Score geral: {score:.2f}\n")
        p.add_run(f"Classificação final: {status}\n\n")

        doc.add_paragraph(
            "Avaliação fundamentada na Política Nacional de Atenção às Urgências (PNAU), "
            "Portarias do SUS e parâmetros utilizados em RAG, PAS e fiscalizações do TCE."
        )

        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer

    if DOCX_DISPONIVEL:
        if st.button("📥 Gerar Relatório DOCX (RAG / PAS / TCE)"):
            arquivo = gerar_relatorio_docx()
            st.download_button(
                "Baixar Relatório",
                data=arquivo,
                file_name="Relatorio_UPA_Dona_Zulmira_RAG_PAS_TCE.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
    else:
        st.info("ℹ️ Para habilitar o DOCX, adicione `python-docx` ao requirements.txt.")

    st.success("✅ Análise concluída com sucesso!")

















